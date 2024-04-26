from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import json
from datetime import datetime
import socket
import platform
import psutil
from datetime import datetime


def read_allure_results(allure_results_path):
    allure_results = []
    for file_name in os.listdir(allure_results_path):
        if file_name.endswith('.json'):
            with open(os.path.join(allure_results_path, file_name), 'r') as file:
                allure_result = json.load(file)
                allure_results.append(allure_result)
    return allure_results

def generate_word_from_allure(feature_name, allure_results):
    # Obtener la fecha y hora actual
    current_datetime = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    # Obtener el nombre del host
    hostname = platform.node()
    # Obtener el nombre de usuario
    username = psutil.users()[0].name
    # Obtener el sistema operativo y su versión
    os_info = platform.platform()
    # Obtener el nombre y versión del sistema operativo
    os_name = platform.system()
    os_version = platform.release()
    # Obtener la información de la CPU
    cpu_info = platform.processor()
    # Obtener el número de núcleos de la CPU
    cpu_cores = psutil.cpu_count(logical=False)
    # Obtener la cantidad total de memoria RAM
    total_memory = psutil.virtual_memory().total
    # Obtener la zona horaria
    timezone_info = datetime.now().astimezone().tzinfo
    # Obtener la fecha y hora actual
    current_datetimex = datetime.now().strftime("%d-%m-%Y - %H:%M:%S")
     
    # Generar el nombre del archivo
    file_name = f"{feature_name}-{current_datetime}.docx"
    
    # Verificar si la carpeta "Evidencias" existe, de lo contrario, crearla
    folder_name = "Evidencias"
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)
    
    # Ruta completa del archivo de Word
    file_path = os.path.join(folder_name, file_name)
    
    doc = Document()
    
    # Agregar título principal
    doc.add_heading('Documento de Evidencia Funcional', level=1)

    # Agregar subtítulo con el nombre del host del equipo del sistema
    doc.add_heading('Detalles de la Ejecucion', level=2)
    doc.add_paragraph(f'Hostname: {hostname}', style='BodyText')
    doc.add_paragraph(f'Usuario: {username}', style='BodyText')
    doc.add_paragraph(f'Sistema Operativo: {os_info}', style='BodyText')
    doc.add_paragraph(f'Procesador: {cpu_info}', style='BodyText')
    doc.add_paragraph(f'Numero de Nucleos de CPU: {cpu_cores}', style='BodyText')
    doc.add_paragraph(f'Memoria RAM Total: {total_memory} bytes', style='BodyText')
    doc.add_paragraph(f'Zona Horaria: {timezone_info}', style='BodyText')
    doc.add_paragraph(f'Fecha y Hora de la Ejecucion: {current_datetimex}', style='BodyText')

    

    # Iterar sobre los resultados de Allure
    for result in allure_results:
        test_case_name = result['name']
        test_case_status = result['status']
        
        # Agregar nombre del caso de prueba y estado
        doc.add_paragraph(f'{test_case_name}\nResultado: {test_case_status}', style='Heading 2')
        
        # Agregar detalles del paso a paso
        steps = result.get('steps', [])
        for step in steps:
            step_name = step.get('name', '')
            step_status = step.get('status', '')
            doc.add_paragraph(f"Paso: {step_name} - Estado: {step_status}", style='BodyText')

            # Agregar imágenes si están disponibles
            attachments = step.get('attachments', [])
            for attachment in attachments:
                if attachment['type'] == 'image/png':
                    image_filename = attachment['name']
                    image_path = os.path.join("./screenshots", image_filename)
                    try:
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run()
                        run.add_picture(image_path, width=Inches(6))  # Ancho de imagen ajustado
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar la imagen
                    except FileNotFoundError:
                        print(f"No se pudo encontrar el archivo: {image_path}")

    # Guardar el documento en la carpeta "Evidencias"
    doc.save(file_path)

# Llamada a la función
allure_results_path = './allure-results'
allure_results = read_allure_results(allure_results_path)
generate_word_from_allure("Evidencia", allure_results)
